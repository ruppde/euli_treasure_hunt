#!/usr/bin/env python3 

# Euli treasure hunt
# Copyright 2020 Arnim Rupp

#    This program is free software: you can redistribute it and/or modify
#    it under the terms of the GNU General Public License as published by
#    the Free Software Foundation, either version 3 of the License, or
#    (at your option) any later version.
#
#    This program is distributed in the hope that it will be useful,
#    but WITHOUT ANY WARRANTY; without even the implied warranty of
#    MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
#    GNU General Public License for more details.
#
#    You should have received a copy of the GNU General Public License
#    along with this program.  If not, see <https://www.gnu.org/licenses/>.


from __future__ import unicode_literals
import sys
import random
import argparse
import os
import shutil
from time import gmtime, strftime
from collections import defaultdict

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_COLOR_INDEX

# install crypto_puzzles in this directory using: git clone https://github.com/2d4d/crypto_puzzles.git
sys.path.append('./crypto_puzzles')
sys.path.append('./examples_cfg')
import crypto_puzzles



################################################################################################
# TODO: get everything between the ########-lines from webpage/django (or at least have it editable there)
################################################################################################

# todo: dirty hack to import example-config passed after -e before arg-parsing in main(), fix with django + db someday
if len(sys.argv) > 1 and sys.argv[1] == '-e':
    exec("from %s import *" % sys.argv[2])
    print("DOING config file: " + sys.argv[2])
else:

    # todo: use some i18n framework like https://github.com/WeblateOrg/weblate or https://github.com/python-babel/babel to replace all the if/elif translations in the code with some
    # kind of message catalog for all languages
    global language
    language='en'
    #language='de'

    # sequence of puzzles in format "puzzle:grade adjustment" (to start with easier puzzles to avoid frustration in the beginning). after they've invested some time, players should be more
    # willing to also crack the next one :)
    puzzle_sequence = """crypto_puzzles/wrong_whitespace:-2
qr_puzzle:0
crypto_puzzles/stego_acrostic:-1
qr_in_qr:0
black_on_black_docx:0
crypto_puzzles/figlet:0
generate_crackme_python:0
substitute_partly_solved_frequency_analysis:0
"""
    # todo: make uv_lamps optional/configurable in sequence and not just in the end


    # players names and school grades (decrease by 1 if the kids have never played it before, increase by 1 or 2 above actual grade if kids have enough experience with this kind of puzzles. In doubt start low to avoid frustration, this systems isn't yet tested and balanced. Feedback welcome ;)
    players_grade={ 'Player-1':8, 'Player-2':3 }
    #players_grade={ 'Alice':7, 'Bob':4 }
    #players_grade={ 'Spieler-1':3,'Spieler-2':5, 'Spieler-3':4, 'Spieler-4':7, 'Spieler-5':1 }

    #players_skills={ }
    players_skills={ 'Player-1':'py' }
    #players_skills={ 'Alice':'py' }

    # remember stuff to avoid double usage of crypto functions and hiding places (per location!)
    player_used_function = defaultdict(str)
    player_used_place = defaultdict(str)

    # Groups of players in the same location. Everybody not mentioned here is alone in his location
    same_location=[ ['Player-1', 'Player-2'] ]
    #same_location=[ ['Alice','Bob'] ]
    #same_location=[ ['Spieler-1','Spieler-2'], ['Spieler-3','Spieler-4'] ]

    # also include plaintext for this players in docx. useful if you want to do manual encryption in the docx or use different kinds of puzzles. don't forget to remove later
    # other method is using -d on the command line
    include_plaintext_for_players=[]
    #include_plaintext_for_players=['Spieler-1', 'Spieler-4']


    # method used to bring all players into the same track at the same time (at the moment only in the end, todo). this can be either:
    # tokens : all players collect tokens, enter them in a webpage and if (nearly) all of them are found, the webpage shows the place of the next stage. only method so far usable in distributed locations for e.g. videochat with all players
    #join_method = 'tokens'

    # combination_lock : every player finds one digit of a combination lock and together they can open it. only works if all locks in all locations use the same combination (todo)
    #join_method = 'combination_lock'

    # join_puzzle: give a sheet with an encrypted clue where several parameters are needed the decrypt it, e.g. vigenere cipher (boring work for longer messages unless some tool is used which again is boring) or a sheet full of characters where the message starts at certain coordinates in a certain direction with every n-th character. would be nice to have something for people, which don't have a combination lock at hand.
    # (doesn't work for players in multiple locations, use tokens)
    # todo: find a better kind if metapuzzle for this purpose
    join_method = 'join_puzzle'

    # none: single player
    #join_method = 'none'
    # btw: spliting players into several tracks to have different puzzles for different age groups/skill sets is easy: just put multiple named envelopes with messages in one place

    # prefix on every token to be found to make it easier to recognize the right number as being part of the game
    # (tokens are like flags in CTFs)
    token_prefix = "Treasure-"
    #token_prefix = "Schatz-"
    # number of digits a token has
    token_len = 4
    # remeber all generated tokens to avoid duplicates
    token_list=[]
    # name of the script where the players enter the found tokens
    token_script = 'treasure.py'
    #token_script = 'schatz.py'

    # the number of digits of the combination lock should be >= number of players in the location
    # (can also be charaters on certain locks or a cryptex)
    combination_lock_combination='066'
    # give a final puzzle for the combination of the combination lock (not possible if join_method = 'combination_lock')
    # choose "round" looking digits for combination_lock_combination like 066, looks more natural ;)
    combination_lock_puzzle='stego_saurus'



    text_known_plaintext={}
    text_known_plaintext['en'] = "Top secret!"
    text_known_plaintext['de'] = "Streng geheim!"

    # where to hide the initial note. should be something personal if several players are in the same place, but 2 notes could also be in one place
    initial_hiding_place = 'under your kitchen sink'
    #initial_hiding_place = 'unter Eurer Küchenspüle'
    # you can also hide a locked treasure chest there and make the players find the combination to open it.
    initial_hiding_locked_chest = True

    # place to hide ultraviolet lamp(s)
    uv_lamps_place = 'in the largest cooking pot'
    #uv_lamps_place = 'in Eurem größten Kochtopf'

    # this can be either the final place of the treasure or e.g. of the combination which opens a treasure chest, the players already have
    final_treasure_place = 'near your vacuum cleaner'
    #final_treasure_place = 'bei Eurem Staubsauger'

    # for all players: block hiding places of treasure and uv-lamps for tokens and stages
    # TODO: lousy coding, relies on 100% identical strings, fix with unique ids with django/database
    for player in players_grade:
        player_used_place[player] += initial_hiding_place + ','
        if uv_lamps_place:
            player_used_place[player] += uv_lamps_place + ','

    token_text = ''
    if join_method =='tokens':
        if language == 'en':
            token_text = '- The secret tokens you find throughout the hunt have to be entered on http://example.com/treasure.py . Once all players have entered all tokens, you will get to know the location of the next to last stage.'
        elif language == 'de':
            token_text = '- Die Geheim-Codes, die Ihr später findet, müsst Ihr auf http://beispiel.de/schatz.py eingeben. Wenn Ihr alle Codes gefunden habt, erfahrt Ihr auf dieser Webseite den Ort der vorletzten Station.'

    players_introduction={}
    players_introduction['en']="""Welcome to the big video chat treasure hunt on the occasion of Jane Does birthday!

The evil spies have stolen all your sweets!!!! Hrngrmblfx!! Cruel!!! But you are lucky, that these morons needed to take notes where they have hidden them. One of our agents could get hold of these notes, below is the first part for you.

However, before you can start, here are some tips:
- The spies start all of their secret clues with the text \"""" + text_known_plaintext['en'] + """\" If you have figured out, how these words are encrypted, you can decrypt the rest of the message.
- If you are stuck, ask the other players, if they have any ideas.
""" + token_text + """

Have a good time!
"""

    players_introduction['de']="""Herzlich willkommen zur großen Video-Chat-Schatzsuche anlässlich von Erika Mustermanns Geburtstags!

Die bösen Spione habe Eure Süßigkeiten gemoppst!!!! Hrngrmblfx!! Gemein!!!1!! Aber zum Glück mussten die Dösbaddel sich selber aufschreiben, wo sie versteckt sind. Diese Aufzeichnungen konnte einer unserer Agenten unter großen Gefahren an sich bringen. Weiter unten ist der erste Teil der Botschaft für Dich.

Doch bevor es losgeht noch ein paar Tipps:
- Die Spione beginnen alle ihre Geheim-Botschaften mit dem Text \"""" + text_known_plaintext['de'] + """\" Wenn man erkannt hat, wie dieser Anfang verschlüsselt ist, kann man auch den Rest der Botschaft entziffern.
- Wenn Ihr nicht weiterkommt, fragt ruhig die Anderen, ob sie eine Lösungsidee haben. 
""" + token_text + """
Und nun viel Erfolg!
"""

    parents_introduction={}
    parents_introduction['en']="""Welcome to the big video chat treasure hunt on the occasion of Jane Does birthday!

To prepare the treasure hunt, you need to hide several secret clues and a share of the treasure for your kids. Place use the table below.

Some tips on hiding:
    - You can put each message in an envelope. That's about as useful as wrapping birthday presents: It's still the same present but kind of fun to unwrap it ;)
    - Don't hide the stuff too well, it should be easily found if someone is looking in the right place. But the players also shouldn't randomly stumble upon it.
    - Don't get caught by the kids during hiding and also dispose these notes, because they contain parts of the solution.
    - If the proposed hiding place is not explicit, because e.g. you have two sinks in the bathroom or your kid does not have a favorite lunch box, just use one of them and they'll have to be searched.
    - Do not fold QR codes, makes it harder for the device to read them.
    - Some puzzles contain a hint on the bottom of the page. You can choose to fold up to cover the hint and stick it with glue or tape. This way players don't see the hint immediatly and have a choice, if they want to open it.
"""

    parents_introduction['de']="""Herzlich willkommen zur großen Video-Chat-Schatzsuche anlässlich von Erika Mustermanns Geburtstags!

Zur Vorbereitung der Schatzsuche müssen einige Geheim-Botschaften und ein Anteil vom Schatz vorher versteckt werden. Dazu bitte die Tabelle unten durchgehen.

Noch einige Tipps zum Verstecken:
    - Die einzelnen Nachrichten können in Umschläge gesteckt werden. Das ist ähnlich sinnvoll wie Geburtstagsgeschenke zu verpacken: Es ist immer noch das gleiche Geschenk, macht aber irgendwie Spaß beim auspacken ;)
    - Nicht zu gut verstecken, man sollte es recht schnell sehen, wenn man an der richtigen Stelle sucht, aber auch nicht zufällig drüber stolpern.
    - Natürlich beim Verstecken nicht von den Kindern erwischen lassen und die übrigen Zettel danach gut entsorgen, weil sie Teile der Lösungen enthalten.
    - Wenn Ihr das vorgesehene Versteck nicht eindeutig ist, z.B. zwei Badwaschbecken habt oder es keine explizite Lieblingsfrühstücksdose gibt, nehmt einfach eins davon und dann muss halt alles abgesucht werden.
    - QR codes nicht falten, erschwert dem Gerät das Lesen.
    - Einige Rätsel haben einen Tipp unten drunter. Bei Bedarf kann man die Seite hochfalten um den Tipp zu verdecken und sie so zukleben. Auf diese Weise sehen die Spieler den Tipp nicht sofort und habe die Wahl, ob sie ihn öffnen wollen.
"""

    text_secret_message={}
    text_secret_message['de'] = 'Eine Geheim-Botschaft für Dich:'
    text_secret_message['en'] = 'A secret message for you:'

    hiding_instructions = defaultdict(str)
    token_cut_outs = defaultdict(str)



    # keep track which hiding and token places in which locations have been used already 
    used_hiding_place_by_location={}
    used_token_place_by_location={}
    universal_hiding_places=[]
    universal_token_places=[]

    # places where something tiny like a small note with a token number can be attached
    # don't use anything in or near universal_hiding_places_text to avoid accidential findings!
    universal_token_places_text={}
    universal_token_places_text['en']="""On the handle of your baking oven
Under the sink where you brush your teeth
Near your toothbrush
On the largest knive in your kitchen
Under your big forks
Under your big spoons
On the window handle in your kitchen
On the backside of the water tap in your kitchen
On the backside of the water tap in your bathroom
On the lamp next to your bed
On your biggest mirror
"""

    universal_token_places_text['de']="""An Eurem Backofengriff
Unter dem Waschbecken wo Du Zähne putzt
In der Nähe Deiner Zahnbürste
Am allergrößten Messer in Eurer Küche
Unter Euren großen Gabeln
Unter Euren großen Esslöffeln
Am Fenstergriff in Eurer Küche
Auf der Rückseite vom Wasserhahn in Eurer Küche
Auf der Rückseite vom Wasserhahn in Eurem Bad
An der Lampe an Deinem Bett
An Eurem größten Spiegel
"""


    # TODO: unused, offer in webgui
    # these places might not be in every flat or house because e.g. the kitchen is integrated into the living room and doesn't have a door 
    # You can reuse existing numbers, like the article number of the noodles, that date on the monument in the park minus 673, a phone number on a billboard, on a webpage, a number in a movie, song, boardgame, book, toy, file on usb-stick, written somewhere with UV-ink, ...
    universal_token_places_ideas_text={}
    universal_token_places_ideas_text['de']="""am Griff der Küchentür
am Griff der Wohnzimmertür
an der Spülmaschine
in einer Taschentuchpackung auf Deiner Fensterbank
"""
    #Fäden zum rausziehen???

    # places where something like an envelope can be hidden
    universal_hiding_places_text={}
    universal_hiding_places_text['en']="""under your dining table
under your sofa cushions
under your soup plates
under your small plates
under your big plates
in your fridge door
in your favorite lunch box
under your chair at the dining table
in the pockets of your warmest jacket
in your wardrobe under to your underpants
in your wardrobe under to your pyjama
in your schoolbag
under your desk
in your letterbox
near your vacuum cleaner
under your kitchen sink
in your largest cooking pot
in your umbrella
under your breadbox/breadbasket
"""

    universal_hiding_places_text['de']="""unter dem Esstisch
unter Euren Sofakissen
unter Euren Suppentellern
unter Euren kleinen Tellern
unter Euren großen Tellern
in Eurer Kühlschranktür
in Deiner Lieblingsfrühstückdose
unter Deinem Stuhl am Essplatz
in den Taschen Deiner dicksten, wärmsten Jacke
in Deinem Kleiderschrank bei den Unterhosen
in Deinem Kleiderschrank bei den Schlafanzügen
in Deinem Schulranzen
unter Deinem Schreibtisch
in Eurem Briefkasten
bei Eurem Staubsauger
unter Eurer Küchenspüle
in Eurem größten Kochtopf
in Deinem Regenschirm
unter Eurem Brotkorb/Brotdose
"""

    # TODO unused, offer in webgui
    # these places might not work in every flat or house:
    hiding_place_ideas="""Unter dem Sofa
hinter der Heizung im Wohnzimmer
Im xxx-Pokal
in Buch xxx
in Brettspiel xxx
im CD-Player
mit Schnur von Schrank ziehen?
unter den Klopapierrollen
unter den Äpfeln
unterm Wohnzimmer-Teppich
im Backofen (Obacht ;)
bei der Mikrowelle (nicht drin ;)
unter Deinem Bett
hinter Bild XY im Wohnzimmer
an Deiner Pinwand
in Deinem Kleiderschrank bei den Sportsachen
"""

#########################################
# bonus, not in the example files:
#########################################

parents_introduction_hide_own_puzzles = {}
parents_introduction_hide_own_puzzles['en'] = '''Including your own puzzles is easy:
- If it's some kind of puzzle box, just put any message inside it and place it, where the message was supposed to go.
- If it can point to another hiding place, just put it between two other messages. Let's say message 3 points to "your largest cooking pot", where message 4 would normally go. Now just make up another place to hide, e.g. "under the piano". Just put message 4 under the piano, make your own puzzle point there and put it in the largest cooking pot.
- Examples of puzzles you can create on your:
    - Take a picture of a very small but distinct detail in your place, e.g. a section of a picture on the wall or a part of the doghouse. Hide the message there and print the photo as the clue.
    - Get a wireless door bell and hide the button in one place and the bell in another. (Test the range!)
    - Change something in your living room, e.g. flip the calendar to the next month or remove a vase. Hide the message there, take a photo of the whole living room, reverse the change and give the printed photo as the clue.
    - ... there are lots of ideas in the web.
'''
parents_introduction_hide_own_puzzles['de'] = '''Eigene Rätsel einfügen ist einfach:
- Wenn es eine Rätsel-Dose ist, einfach eine Nachricht reinstecken und sie dort verstecken, wo die Nachricht landen sollte.
- Wenn es zu einem anderen Versteck zeigen kann, einfach zwischen zwei andere Nachrichten einfügen. Nehmen wir an Nachricht 3 zeigt zu "dem größten Kochtopf", in dem dann normalweise Nachricht 4 kommen würde. Dann einfach ein neues Versteck ausdenken, z.B. "unter dem Klavier". Dann Nachricht 4 unter dem Klavier verstecken, das eigene Rätsel dorthin zeigen lassen und in den größten Kochtopf stecken.
- Beispiele für eigene Rätsel:
    - Mache ein Foto von einem kleinen aber eindeutigen Detail der Wohnung, z.B. ein Ausschnitt von einem Bild an der Wand oder einem Teil der Hundehütte. Verstecke die Nachricht dort und nehme einen Ausdruck des Fotos als Hinweis.
    - Besorge eine drahtlose Türklingel und verstecke den Klingelknopf an der einen Stelle und die Klingel an der anderen. (Reichtweite testen!)
    - Ändere etwas im Wohnzimmer, z.B. den Kalender einen Monat weiterblättern oder entferne eine Vase. Verstecke die Nachricht dort, fotografiere das ganze Wohnzimmer, mach die Änderung rückgängig und verwende einen Ausdruck des Fotos als Hinweis.
    - ... im Internet gibt's jede Menge Ideen.
'''

################## end of stuff for webpage

# distribute digits of combination lock to each player per location
combination_lock_digit_for_player={}
for players_in_location in same_location:
    digits = list(combination_lock_combination[::-1])

    # give player more than one digit of it's a long combination
    num_to_pop = int(len(digits) / len(players_in_location))


    num_player = 0
    for player in players_in_location:
        num_player += 1
        num_digit = 0
        digit_for_player =''
        position = ''
        for lola in range(num_to_pop):
            digit_for_player += digits.pop()
            position += str((num_player-1) * num_to_pop + 1 + num_digit) + ','
            num_digit += 1

        position = position.rstrip(',')

        position_word = crypto_puzzles.convert_num_to_number_words(position, language)

        # convert e.g. 5 => five, so players can't just look for the few numbers in all the text
        digit_for_player_word = crypto_puzzles.convert_num_to_number_words(digit_for_player, language)

        if language == 'en':
            combination_lock_digit_for_player[player] = 'The combination of position(s) ' + position_word + ' is ' + digit_for_player_word + ' (If you are only missing one digit, just try it out ;)'
        elif language == 'de':
            combination_lock_digit_for_player[player] = 'Die Gepäcknummer für Koffer dreiundzwanzig an Position ' + position_word + ' ist ' + digit_for_player_word

        # if only one digit is left over, it can easily be brute forced and that's the best way to prevent players from brute forcing a digit they're supposed to find because then they would have to try 100 digits ;) so we don't give it the last player and leave it like this

# put text above into lists:
for line in universal_hiding_places_text[language].splitlines():
    universal_hiding_places.append(line)

for line in universal_token_places_text[language].splitlines():
    universal_token_places.append(line)


text_the_next_secret_is={}
the_next_secret_is_en_tmp = [' The next secret message is ', ' Another clue is ', ' A further secret message is ', ' The next secret is ', ' If you are looking for another message, it is ']
text_the_next_secret_is['en'] = the_next_secret_is_en_tmp

the_next_secret_is_de_tmp = [' Die nächste Geheim-Botschaft ist ', ' Eine weitere Botschaft ist ', ' Noch ein Hinweis ist ', ' Das nächste Geheimnis ist ', ' Wenn Du noch ein Geheimnis suchst, es ist ']
text_the_next_secret_is['de'] = the_next_secret_is_de_tmp

text_hidden_word={}
hidden_word_en_tmp = ['hidden ', 'concealed ', 'stashed ']
text_hidden_word['en'] = hidden_word_en_tmp

hidden_word_de_tmp = [' versteckt!', ' verborgen!', ' gut verborgen!', ' gut versteckt!']
text_hidden_word['de'] = hidden_word_de_tmp

hint_plaintext_attack={}
hint_plaintext_attack['en'] = 'Remember to look for \'' + text_known_plaintext['en'] + '\''
hint_plaintext_attack['de'] = 'Denk an die Suche nach \'' + text_known_plaintext['de'] + '\''

def init_log(outdir):
    global logfile
    logfile = open(outdir + '/logfile.txt', 'w')

def log(text, level=1):
    if level <= loglevel:
        # log to stdout and file
        print(text)
        logfile.write(text + '\n')


def randomize_hiding_places(players_grade, same_location):
    # create a randomized list for each player, where stuff for the next round is hidden, usually in someone elses flat if people are spread
    location_player_from_to={}

    # todo: avoid the same pairing as in the last round to maximize interactions with different players

    if len(players_grade) == 1:
        # only 1 player, everything will be in the same location
        player=list(players_grade)[0]
        location_player_from_to[player] = player
    else:

        # repeat until perfect draw
        count_draw = 0
        while True:
            players_lottery_wheel = list(players_grade.keys())
            count_draw += 1
            # lousy way of drawing a player and checking if he is another flat until that's the case or it's obviously the last player
            for player_from in players_grade:

                    count = 0
                    while True:
                        count += 1
                        players_in_same_location=0

                        player_to_pos = random.randrange(len(players_lottery_wheel))
                        player_to = players_lottery_wheel[ player_to_pos ]

                        if player_from != player_to:
                            # check for different location
                            for players_in_location in same_location:
                                hits = 0
                                for player_tmp in players_in_location:
                                    if player_tmp == player_to:
                                        hits += 1
                                    elif player_tmp == player_from:
                                        hits += 1
                                if hits == 2:
                                    players_in_same_location=1

                        # take this roll if players are in different location or enough attempts have been made to avoid this
                        if player_from != player_to and not players_in_same_location or count > 20:
                            player_to = players_lottery_wheel.pop(player_to_pos)
                            log("from, to: " + player_from + player_to, 8)
                            location_player_from_to[player_from] = player_to
                            break

            # sometimes in the last draw of player_to there was only player_from left, repeat in this case
            if player_from != player_to:
                if not players_in_same_location or count_draw > 20:
                    log('found pairing:' +  player_from + ':' +  player_to + ':' +  str(players_in_same_location) + ':' + str(count_draw))
                    break

    return location_player_from_to


def get_all_players_in_this_location(player):

    # at least set the one passed player
    all_players_in_this_location=[player]

    # check for more
    for players_in_location in same_location:
        if player in players_in_location:
            all_players_in_this_location=players_in_location
            break

    return all_players_in_this_location

def generate_token(player):
    while True:
        token = token_prefix + str(random.randrange(10**(token_len-1), 10**token_len-1))
        # remeber all generated tokens to avoid duplicates
        if token not in token_list:
            token_list.append(token)
            break

    # save generated tokens for later use, todo: use database
    tokenfile.write(token + '\n')

    return token
    
def init_random_stuff_per_round():
        log("init_random_stuff_per_round")
        # generate random list of which players gets the hints for which other player, where his token is hidden
        location_player_from_to = randomize_hiding_places(players_grade, same_location)
        log("location_player_from_to:" + str(location_player_from_to))

        # generate random order of places, token places and functions
        hiding_place_lottery_wheel = universal_hiding_places.copy()
        random.shuffle(hiding_place_lottery_wheel)

        token_place_lottery_wheel = universal_token_places.copy()
        random.shuffle(token_place_lottery_wheel)

        crypto_functions_lottery_wheel = crypto_puzzles.get_crypto_functions()
        random.shuffle(crypto_functions_lottery_wheel)

        # multiply the lists to be able to handle more players than crypto functions
        for x in range(20):
            hiding_place_lottery_wheel.extend(hiding_place_lottery_wheel)
            token_place_lottery_wheel.extend(token_place_lottery_wheel)
            crypto_functions_lottery_wheel.extend(crypto_functions_lottery_wheel)

        return location_player_from_to, hiding_place_lottery_wheel, crypto_functions_lottery_wheel, token_place_lottery_wheel

def make_rows_bold(*rows):
    # for docx
    for row in rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

def get_random_crypto_function(function_type, crypto_functions_lottery_wheel, player):
    # get a random function, which wasn't used before for this player, to encrypt the place of the next stage
    # todo: this code sucks, fix someday with more time ...

    if function_type == 'reversible':
        all_reversible_crypto_functions = get_crypto_functions('reversible')

    while True:
        crypto_function = crypto_functions_lottery_wheel.pop()
        if player in player_used_function and crypto_function in player_used_function[player]:
            log("redraw function, already used")
        else:
            if function_type !='reversible' or ( function_type == 'reversible' and crypto_function in all_reversible_crypto_functions ):
                # remember, that this player already used this function
                player_used_function[player] += crypto_function + ','
                break
    return crypto_function

def docx_change_font_size(table, size_odd, size_even):
    count=0
    for row in table.rows:
        count+=1
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font

                    if count /2 == int(count/2):
                        font.size= Pt(size_even)
                    else:
                        font.size= Pt(size_odd)

def format_docx(docx):
    # set font style
    tstyle = docx.styles['Normal']
    font = tstyle.font
    font.name = "Calibri" # works on Linux, don't know about win (todo?)
    font.size = Pt(9)
    return docx

def text_hide_next_round( round_num, total_rounds, player, hiding_place, language):
    # todo, function called only once at moment, remove??
    if round_num < total_rounds or join_method == 'join_puzzle':
        if language == 'en':
            hiding_instructions = ("Message " + str(round_num + 1) + ' for ' + player + ':::' + hiding_place + '\n')
        elif language == 'de':
            hiding_instructions = ("Nachricht " + str(round_num + 1) + ' für ' + player + ':::' + hiding_place + '\n')
    # last round, nothing to hide anymore
    else:
        hiding_instructions = ""

    return hiding_instructions

def generate_hiding_message(hiding_place, language):
    if language == 'en':
        # build e.g.: Top secret! + The next secret is + hidden + under your desk
        message = text_known_plaintext[language] + text_the_next_secret_is[language][random.randrange(len(text_the_next_secret_is[language]))] + text_hidden_word[language][random.randrange(len(text_hidden_word[language]))] + hiding_place 

    elif language == 'de':
        # build German version of e.g.: Top secret! + The next secret is + hidden + under your desk
        message = text_known_plaintext[language] + text_the_next_secret_is[language][random.randrange(len(text_the_next_secret_is[language]))] + hiding_place + text_hidden_word[language][random.randrange(len(text_hidden_word[language]))] 

    return message

def hiding_instructions_file(round_num, crackme_short_name, language):
    if language == 'en':
        return ('Follow instructions in message ' + round_num + ' and put file ' + crackme_short_name + ' on e.g. a USB-stick with the message into an envelope:::See below\n')
    elif language == 'de':
        return ('Beachte Anleitung in Nachricht ' + round_num + ' und lege Datei ' + crackme_short_name + ' auf z.B. einen USB-Stick mit der Nachricht in einen Umschlag:::Siehe unten\n')


def hiding_instructions_file_message(docx, crackme_short_name, language):
    if language =='en':
        docx.add_paragraph('✁ - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - ')
        docx.add_heading('INSTRUCTIONS to prepare this message, cut off')
        docx.add_paragraph('Put the file ' + crackme_short_name + ' from this zip archive onto a USB-stick (or any other medium the players can read, e.g. SD-Card (Compact Flash, Memory Stick, ...), CD-R, floppy, QIC80-tape, MP3-player, the ebook reader with the broken display which still works as mass storage device, ...) and with the message above into an envelope (or just stick them together)' +  '\n')

    elif language =='de':
        docx.add_paragraph('✁ - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - ')
        docx.add_heading('ANLEITUNG zur Vorbereitung dieser Nachricht, abschneiden')
        docx.add_paragraph('Lege Datei ' + crackme_short_name + ' aus diesem ZIP-Archiv auf einen USB-Stick (oder irgendein anderes Medium, dass die Spieler lesen können: Webseite, SD-Card (Compact Flash, Memory Stick, ...), CD-R, Floppy, Quietschebändchen, MP3-Player, der eBook-Reader mit dem kaputten Display der sich aber noch als Massenspeicher meldet, ...) mit der Nachricht in einen Umschlag (oder alternativ einfach zusammenkleben).' +  '\n')
    return docx

###################################################################################
def main():

    # parse command line args
    parser = argparse.ArgumentParser()
    parser.add_argument("--loglevel", "-l", help="set log level", default=1)
    parser.add_argument("--example", "-e", help="Generate example")
    parser.add_argument("--debug", "-d", action='store_true', help="Debugging, include plaintext in docx")
    args = parser.parse_args()
    global loglevel
    loglevel = int(args.loglevel)
    global debug
    debug = int(args.debug)
    example = args.example

    outdir=""
    if example:
        outdir += './examples/'
        outdir += 'euli_' + example + '_'
    else:
        outdir += 'euli_' + strftime("%Y%m%d-%H%M%S", gmtime()) + '_'

    for player in players_grade:
        outdir += player + '_' + str(players_grade[player]) + '_'

    if join_method =='join_puzzle':
        # generate each players individual message for the join_puzzle
        join_puzzle_complete_message = generate_hiding_message(uv_lamps_place, language)
        lowest_grade = min(list(players_grade.values()))
        player_names = list(players_grade.keys()) 
        join_puzzle_messages = crypto_puzzles.join_puzzle(join_puzzle_complete_message, language, lowest_grade, player_names)


    # save docx with the config for this game plus one line for putting it into an .md table overview with all generated example games
    docx = Document()
    docx = format_docx(docx)
    md_requirements=''
    md_players=''

    outdir += '_NEEDS:'

    docx.add_heading('Requirements to play this game')
    if 'qr_puzzle' in puzzle_sequence or 'qr_in_qr' in puzzle_sequence:
        docx.add_paragraph('In each location: Device which can read QR codes')
        outdir += 'qr-reader'
        md_requirements +='QR-reader, '

    if 'generate_crackme_python' in puzzle_sequence:
        docx.add_paragraph('In each location: A player with Python skills, a device which can read read an USB stick/SD-card and execute Python scripts.')
        outdir += '_PC_and_python-skills'
        md_requirements+='Python skills, '

    if 'black_on_black_docx' in puzzle_sequence:
        docx.add_paragraph('In each location: A player with basic computer skills, a device which can read read an USB stick/SD-card and contains a word processor.')
        outdir += '_PC_with_word_processor'
        md_requirements+='PC with word processor, '

    if uv_lamps_place:
        docx.add_paragraph('In each location: UV pocket lamp (one per player in the location would be perfect) and UV marker')
        outdir += '_uv-lamps'
        md_requirements+='UV marker + UV pocket lamp(s), '

    if initial_hiding_locked_chest:
        docx.add_paragraph('In each location: Some kind treasure chest with combination lock.')
        outdir += '_chest'
        md_requirements+='treasure chest with combination lock, '

    if 'example' in outdir:
        shutil.rmtree(outdir, ignore_errors=True, onerror=None)
    os.makedirs(outdir)
    init_log(outdir)

    docx.add_heading('Settings used to generate this game')
    docx.add_paragraph('Starting place (tell the players to look here for the first secret note once the game starts): ' + initial_hiding_place)
    docx.add_paragraph('Language (ISO code): ' + language)
    docx.add_paragraph('Known plaintext used at begining of every secret message: ' + text_known_plaintext[language])

    if join_method == 'tokens':
        global tokenfile
        tokenfile = open(outdir + '/tokens.txt', 'w')

        docx.add_paragraph('Tokens activated: Each secret note contains a hint to a place in the other players locations, where a token is hidden. Players have to cooperate to collect all tokens. They have to entered into a webpage to get the pre-final location on the way to the treasure.')
        docx.add_paragraph('Token prefix: ' + token_prefix )

    docx.add_heading('Players and school grades')
    for player in players_grade:
        docx.add_paragraph(player + '  grade: ' + str(players_grade[player]))
        md_players += player + ' (' + str(players_grade[player]) + '), '

    docx.add_heading('Players in same location (to avoid using the same hiding place twice)')
    if same_location:
        for players_in_location in same_location:
            docx.add_paragraph(', '.join(players_in_location))
    else:
        docx.add_paragraph('None')

    docx.add_heading('Sequence of puzzles : Grade adjustment (usually to make it easier in the beginning)')
    docx.add_paragraph(puzzle_sequence)

    if join_method =='tokens':
        docx.add_heading('Universal hiding places for tokens used')
        docx.add_paragraph(universal_token_places_text[language])

    docx.add_heading('Universal hiding places for stages used')
    docx.add_paragraph(universal_hiding_places_text[language])

    filename = outdir + '/requirements_and_settings.docx'
    docx.save(filename) 
    log('Save config document: ' + filename)
    # done writing config docx

    # write md line:
    md_filename = outdir + '.zip'
    # fix directory for this usage:
    md_filename = md_filename.replace('/examples/', '/examples_zip/')

    md_players = md_players.rstrip(', ')
    md_requirements = md_requirements.rstrip(', ')
    md_table_file = open(outdir + '/overview_table_line.md', 'w')
    # markdown link to local zip
    md_table_file.write('[zip](' + md_filename + ') | ' + md_players + ' | ' + md_requirements + ' | ' + language + '\n')
    md_table_file.close()



    # init docx for each player
    docx={}
    for player in players_grade:
        log('Init document for player: ' + player)
        docx[player] = Document()

        docx[player] = format_docx(docx[player])

        numlock_msg=''
        if language == 'en':
            if combination_lock_combination:
                numlock_msg = 'Set the combination lock of the chest to ' + combination_lock_combination + '. '
            # todo: fix: instructions are sorted alphabetically at the moment so this sentence has to start with 'A' to be on top
            hiding_instructions[player] += ('A good start: ' + numlock_msg + 'Put the treasure (whatever is valuable to the players, usually sweets ;) into ... :::lockable treasure chest.\n')
        elif language == 'de':
            if combination_lock_combination:
                numlock_msg = 'Stellen Sie das Zahlenschloss der Schatztruhe auf ' + combination_lock_combination + '. '
            # todo: fix: instructions are sorted alphabetically at the moment so this sentence has to start with 'A' to be on top
            hiding_instructions[player] += ('Als erstes: ' + numlock_msg + 'Den Schatz (was auch immer wertvoll für die Spieler ist, üblicherweise Süßigkeiten ;) verstecken in ... :::verschließbarer Schatztruhe.\n')
        else:
            print("unknown language: " + language)
            sys.exit(1)


            if join_method =='tokens':
                # copy script for code collection webserver into outdir
                shutil.copy2('contrib/' + token_script, outdir)



    total_rounds = len(puzzle_sequence.splitlines())
    round_num = 0
    for puzzle_and_grade_adj in puzzle_sequence.splitlines():
        round_num += 1
        log('Generating tasks and puzzles for round ' + str(round_num))

        puzzle, grade_adjustment = puzzle_and_grade_adj.split(':')
        
        
        location_player_from_to, hiding_place_lottery_wheel, crypto_functions_lottery_wheel, token_place_lottery_wheel= init_random_stuff_per_round()

        # create puzzles into seperate .docx for each player
        for player in players_grade:
            log('Doing player ' + player)
            grade = players_grade[player] + int(grade_adjustment)
            hint=""

            if grade < 1: grade = 1 

            log('Doing: ' + player + ' adjusted grade: ' + str(grade) )

            if language == 'en':
                header = 'Message ' + str(round_num ) + ' for ' + player
            elif language == 'de':
                header = 'Nachricht ' + str(round_num ) + ' für ' + player

            docx[player].add_heading(header)

            if round_num == 1:
                if initial_hiding_locked_chest == True:
                    if language == 'en':
                        docx[player].add_paragraph('You should have found a treasure chest with this message! To find out how to open it, read on ;)' )
                    elif language == 'de':
                        docx[player].add_paragraph('Mit dieser Nachricht solltest Du eine Schatztruhe gefunden haben! Um herauszufinden, wie man sie öffnen kann, lies weiter ;)' )

                # Message for use in a window envelope, someday again (todo)
                #if language == 'en':
                    #docx[player].add_paragraph('Do not open before the treasure hunt starts!' )
                #elif language == 'de':
                    #docx[player].add_paragraph('Erst öffnen, wenn die Schatzsuche beginnt!' )
            else:
                docx[player].add_paragraph(' ' )

            # include players_introduction on 1st page
            if round_num == 1:
                docx[player].add_paragraph(players_introduction[language])

            # if this isn't the 1st stage include a token place
            # (the token comes above the puzzle on the page because it's quickly communicated, then the puzzle can be solved)
            else:
                if join_method =='tokens':
                    # pick ANOTHER random player which gets the hint where to find a token in his location
                    player_to = location_player_from_to[player]
                    log('This token gets hidden in the location of player_to: ' + player_to)

                    # an unused random place to hide the token 
                    found=False
                    while not found:
                        token_place = token_place_lottery_wheel.pop()
                        log("token_place:" + token_place)
                        all_players_in_this_location = get_all_players_in_this_location(player)
                        log("all players in this location:" +str(all_players_in_this_location))

                        # check all players in this location
                        already_used=0
                        for player_in_this_location in all_players_in_this_location:
                            if player_in_this_location in player_used_place and token_place in player_used_place[player_in_this_location]:
                                log("redraw token place, already used")
                                already_used=1
                                break

                        if not already_used:
                            player_used_place[player] += token_place + ','
                            # TODO breaks on too many players in one location for too few hiding places
                            found = True
                            break
                    log("token place:" + player + ':' +  token_place)

                    # generate a random token
                    token = generate_token(player)

                    if player != player_to:
                        # if the token is for another player
                        if language == 'en':
                            token_message = 'Well done! Give the following tip to player ' + player_to + ': ' + token_place + ' is a secret token for the treasure! Tell ' + player_to + ' about it because all these secret tokens need to be found to find the treasure!'
                        elif language == 'de':
                            token_message = 'Gut gemacht! Gib den folgenden Tipp an ' + player_to + ' weiter: ' + token_place + ' ist ein Geheim-Code für den Schatz versteckt! Sag ' + player_to + ' Bescheid, denn nur wenn Ihr alle Codes gefunden habt, könnt Ihr gemeinsam den Schatz finden!'
                    else:
                        # only one player (group), token in own location
                        if language == 'en':
                            token_message = 'Well done! ' + token_place + ' is a secret token for the treasure!'
                        elif language == 'de':
                            token_message = 'Gut gemacht! ' + token_place + ' ist ein Geheim-Code für den Schatz versteckt!'

                    log("Token msg:" + token_message)
                    docx[player].add_paragraph(token_message)


                if join_method =='tokens':
                    #  instructions for the parents of the other player to hide the token
                    if language=='en':
                        hiding_instructions[player_to] += ("Small slip with token " + token + ':::' + token_place + '\n')
                    if language=='de':
                        hiding_instructions[player_to] += ("Kleiner Code-Zettel " + token + ':::' + token_place + '\n')

                    token_cut_outs[player_to] += (token + '\n')

            # get an unused random place of the next stage for this player
            found=False
            lola=0
            while not found:
                hiding_place = hiding_place_lottery_wheel.pop()
                all_players_in_this_location = get_all_players_in_this_location(player)
                log("all players in this location:" +str(all_players_in_this_location))

                # check all players in this location
                already_used=0
                for player_in_this_location in all_players_in_this_location:
                    if player_in_this_location in player_used_place and hiding_place in player_used_place[player_in_this_location]:
                        log("redraw hiding place, already used")
                        already_used=1
                        lola += 1
                        if lola == 100:
                            print("ERROR: not enough places in: " + sys.argv[2] + " (Reduce number of puzzles or insert more places)")
                            sys.exit(1)
                        break

                if not already_used:
                    player_used_place[player] += hiding_place + ','
                    found = True
                    break

            log("hiding place:" + hiding_place + ':' + player)

            message = ""
            # if not the last round ...
            # build message like "Top secret! The next secret is hidden at XXX"
            if round_num < total_rounds:

                message = generate_hiding_message(hiding_place, language)

            else:
                # last round, no more preparing stuff for the next round
                if join_method =='tokens':
                    # in last round don't point to next stage but give the last token
                    token = generate_token(player)
                    token = crypto_puzzles.convert_num_to_number_words(token, language)

                    if language == 'en':
                        message = text_known_plaintext[language] + ' And finally one more secret token directly for you: ' + token + '. '
                    elif language == 'de':
                        message = text_known_plaintext[language] + ' Und zum Abschluss noch ein Geheim-Code direkt für Dich: ' + token + '. '

                elif join_method =='combination_lock':
                    # hand out combination for the lock
                    message = text_known_plaintext[language] + ' ' + combination_lock_digit_for_player[player]

                elif join_method =='join_puzzle':
                    # join_puzzle is different than combination_lock and tokens in the way, that the last puzzle doesn't give the last information to find the final stage, but the 
                    # location of the note with the join_puzzle. so we need a normal hiding_message as for any other round and after processing this, we add the join_puzzle
                    message = generate_hiding_message(hiding_place, language)

                elif join_method =='none':
                    # no join because single player, just give the location of the uv lamps in the last puzzle
                    message = generate_hiding_message(uv_lamps_place, language)

            log('message:' + message)

            # where to hide the initial note
            if round_num == 1:
                if initial_hiding_locked_chest == True:
                    if language =='en':
                        hiding_instructions[player] += ('Message ' + str(round_num) + ' for ' + player + ' together with the treasure chest (do not hide too good, it is locked anyway ;)'  + ':::' + initial_hiding_place + '\n')
                    elif language =='de':
                        hiding_instructions[player] += ('Nachricht ' + str(round_num) + ' für ' + player + ' zusammen mit Schatztruhe (nicht groß verstecken, ist ja abgeschlossen ;)'  + ':::' + initial_hiding_place + '\n')
                else:
                    if language =='en':
                        hiding_instructions[player] += ("Message " + str(round_num) + ' for ' + player + ':::' + initial_hiding_place + '\n')
                    elif language =='de':
                        hiding_instructions[player] += ("Nachricht " + str(round_num) + ' für ' + player + ':::' + initial_hiding_place + '\n')


            # now a generate a puzzle for this player which tells, where the next stage is hidden

            if puzzle == 'generate_crackme_python':

                # check if player has python skills
                if player in players_skills and 'py' in players_skills[ player ]:
                    # player knows python, go on
                    crackme_code, hint = crypto_puzzles.generate_crackme_python(message, language, grade, 0 )

                    crackme_short_name = 'crackme_for_player_' + player + '.py'
                    crackme_filename = outdir + '/' + crackme_short_name
                    crackme_file = open(crackme_filename, 'w')
                    crackme_file.write(crackme_code)
                    crackme_file.close()

                    docx[player].add_heading(text_secret_message[language])
                    if language == 'en':
                        docx[player].add_paragraph('On this medium you will find a small program, which contains a secret. Use a Python interpreter of your choice to run it. You might need to edit it to reveal the secret ;)')

                    elif language == 'de':
                        docx[player].add_paragraph('Auf diesem Medium wirst Du ein kleines Programm finden, das eine verborgene Nachricht enthält. Benutze einen Python Interpreter Deiner Wahl um es auszuführen. Kann sein, dass Du es ändern musst, um ihm sein Geheimnis zu entlocken ;)')

                    docx[player] = hiding_instructions_file_message(docx[player], crackme_short_name, language)

                    hiding_instructions[player] += hiding_instructions_file(str(round_num), crackme_short_name, language)
                    
                else:
                    # no py skills, fallback to random crypto_puzzles
                    puzzle = 'crypto_puzzles'

            # rot13 is just a special case of a crypto puzzle
            crypto_function = ""
            if puzzle == 'rot13':
                puzzle = 'crypto_puzzles'
                if grade >= 3:
                    crypto_function = "rot13"
                else:
                    # rot13 might be too difficult for grade 1&2, give then some normal puzzle
                    crypto_function = ""

            # don't use elif here because the variable puzzle might have been changed above because a player doesn't have a needed skill for e.g. crackme
            if 'crypto_puzzles' in puzzle:

                # get function out of e.g. crypto_puzzles/whatever
                if "/" in puzzle:
                    dummy, crypto_function = puzzle.split('/')

                if not crypto_function:
                    crypto_function = get_random_crypto_function('all', crypto_functions_lottery_wheel, player)

                log("crypto function:" + str(crypto_function))

                # call function out of variable
                func = getattr(crypto_puzzles, crypto_function)
                encrypted_message, hint = func(message, language, grade )

                if hint:
                    hint += '\n' + hint_plaintext_attack[language]
                else:
                    hint = hint_plaintext_attack[language]

                log('encrypted_message:' + encrypted_message)
                if hint: log('hint:' + hint)

                # save encrypted message to docx
                docx[player].add_heading(text_secret_message[language])

                paragraph= docx[player].add_paragraph()
                font = paragraph.add_run(encrypted_message).font

                if crypto_function == 'figlet':
                    font.name = "FreeMono" # works on Linux, don't know about win (todo?)

                #font.size = Pt(7)

                if hint:
                    if language == 'en':
                        hint = 'Hint: ' + hint
                    elif language == 'de':
                        hint = 'Tipp: ' + hint

                docx[player].add_paragraph(hint)


            elif puzzle == 'qr_puzzle':

                # make the message ǝpᴉsdn uʍop. works with all QR reader's I've tested so far because it only changes utf8 chars
                encrypted_message, hint = crypto_puzzles.upside_down(message, language, grade-1)
                log("qr_puzzle_message: " + message)

                qr_filename = outdir + '/qr_puzzle_' + player + '.png'
                crypto_puzzles.qr_code(encrypted_message, language, qr_filename)

                docx[player].add_paragraph('✁ - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - ')
                # todo: increase "2 snips into 4 pieces" by grade?
                # todo: blank out certain parts of QR code, still readable because of redundancy in the code
                if language == 'en':
                    docx[player].add_heading('INSTRUCTIONS to prepare and hide the QR code')
                    docx[player].add_paragraph('The QR code below has to be cut (don\'t rip) with 2 snips into 4 pieces and put into an envelope. Do not fold the pieces, otherwise it might be to difficult to put them back together again in a way, that will be recognized by the QR scanner. ')
                elif language == 'de':
                    docx[player].add_heading('ANLEITUNG zum Vorbereiten und Verstecken des QR-Codes')
                    docx[player].add_paragraph('Den QR-Code unten mit 2 geraden Schnitten (nicht reißen) in 4 Teile teilen und in einen Umschlag stecken. Die Teile nicht knicken, sonst wird das zu schwierig das nach dem Zusammenpuzzlen mit den Lücken zu scannen. ')

                docx[player].add_paragraph('✁ - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - ')
                docx[player].add_picture(qr_filename)
                os.remove(qr_filename)

                # instructions for the parents of this player to cut the QR code (this round)
                if language == 'en':
                    hiding_instructions[player] += ('Follow instructions in message ' + str(round_num) + ' and put the cut QR code with the message into an envelope:::See below\n')
                elif language == 'de':
                    hiding_instructions[player] += ('Anleitung in Nachricht ' + str(round_num) + ' befolgen und den QR-Code zerschnitten in Umschlag stecken mit der Nachricht:::Siehe unten\n')


            elif puzzle == 'qr_in_qr':

                qr_filename = outdir + '/qr_in_qr_' + player + '.png'
                crypto_puzzles.qr_inside_qr('👍 ' + message, '', language, grade, qr_filename)

                if language == 'en':
                    docx[player].add_heading('2 in 1')
                    docx[player].add_paragraph('Hmmm, this QR code looks somehow different... 🤔🧐')
                elif language == 'de':
                    docx[player].add_heading('2 in 1')
                    docx[player].add_paragraph('Hmmm, dieser QR-Code sieht irgendwie anders aus ... 🤔🧐')

                docx[player].add_picture(qr_filename)
                os.remove(qr_filename)

                docx[player].add_paragraph(hint)

            elif puzzle == 'black_on_black_docx':

                # create docx with secret message in black on black background (hyper secure encryption last seen in real life in 2019 for a new firewall password ;)

                # the code writes alternatly into 2 docx: 
                # 1. black_docx for the puzzle 
                # 2. docx[player] for hidding instructions


                docx[player].add_heading(text_secret_message[language])

                # if grade < 3 leave some letters visible before the black background to give a hint
                if grade < 3:
                    white_message = message[0:6]
                    message = message[6:]
                elif grade < 5:
                    white_message = message[0:2]
                    message = message[2:]
                elif grade < 7:
                    white_message = message[0:1]
                    message = message[1:]
                else:
                    # add some easy crypto for higher grades to slow them down a bit :)
                    message, hint = crypto_puzzles.shift_words(message, language, 5)
                    white_message = ""

                black_docx = Document()
                if language == 'en':
                    docx_short_name = 'message_for_player_' + player + '.docx'
                    black_docx.add_heading('Message for ' + player)
                    black_docx.add_paragraph('This mega most secure of the whole wide world document can not be read by anybody, really nobody, not possible to read not, not a bit, no!!!! Quit, victim!!!!!1!!11one11!!')
                    docx[player].add_paragraph('On this medium you will find a text file, which contains a secret. Use a text processor of your choice to view it. Usually double clicking is enough.')
                elif language == 'de':
                    docx_short_name = 'Nachricht_für_Spieler_' + player + '.docx'
                    black_docx.add_heading('Nachricht für ' + player)
                    black_docx.add_paragraph('Dieses mega-am-sichersten-von-der-ganzen-wo-es-auf-der-Welt-gibt-Dokument kann Niemand, also gar keiner, kann es nicht lesen, überhaupt nicht, nein!!!! Gib auf, Du Opfer!!!1!!!11elf11!!')

                    docx[player].add_paragraph('Auf diesem Medium wirst Du eine Text-Datei finden, die eine verborgene Nachricht enthält. Benutze ein Textverarbeitungsprogramm Deiner Wahl um die Datei zu öffnen. Üblicherweise reicht ein Doppelklick auf die Datei.')

                docx[player] = hiding_instructions_file_message(docx[player], docx_short_name, language)
                hiding_instructions[player] += hiding_instructions_file(str(round_num), docx_short_name, language)

                docx_filename = outdir + '/' + docx_short_name

                paragraph=black_docx.add_paragraph(white_message)
                font = paragraph.add_run(message).font
                font.highlight_color = WD_COLOR_INDEX.BLACK

                black_docx.save(docx_filename)




            elif puzzle == 'substitute_partly_solved_frequency_analysis':

                if language == 'en':
                    message += '. And here a really useless sentence: The ears are left and right, not below!'
                elif language == 'de':
                    message += '. Und hier noch ein sinnloser Satz: Die Ohren sind rechts und links, und nicht unten!'

                # do 123 => one two three to avoid leaving the numbers in plain text
                message = crypto_puzzles.convert_num_to_number_words(message, language)

                encrypted_message, hint, count_letters, num_for_letter, num_to_letter = crypto_puzzles.substitute_partly_solved_frequency_analysis(message, language, grade)

                #docx[player].add_page_break()

                docx[player].add_heading(text_secret_message[language])

                # don't print the nice prefilled table for higher grades
                if grade > 11:
                    docx[player].add_paragraph(encrypted_message)
                else:
                    # create table with cell for each letter

                    # prefill more chars for lower grades
                    if grade > 7:
                        prefill_chars = "e"
                    elif grade > 2:
                        prefill_chars = "ensia"
                    elif grade > 1:
                        prefill_chars = "ensiratuo"
                    else:
                        prefill_chars = "ensiratouchm"

                    log("prefill_chars:" + prefill_chars)

                    secrets = encrypted_message.split(',')

                    num_chars=len(secrets)

                    # Number of columns you want
                    chars_per_row=25
                    num_rows = int(num_chars / chars_per_row) + 1 

                    # create double rows, one for secret, one for plain text solution
                    table = docx[player].add_table(rows=num_rows *2 , cols=chars_per_row)

                    shading_elm={}
                    for row in range(num_rows):
                        for col in range(chars_per_row):

                            position = row * chars_per_row + col
                            if position < len(secrets):
                                secret = secrets[position]

                                cell = table.cell(row *2 , col)
                                cell.text = secret

                                if secret.isdigit():
                                    char = num_to_letter[secret]
                                    if char in prefill_chars:
                                        cell = table.cell(row *2 +1 , col)
                                        cell.text = char
                                else:
                                    if secret == ' ':
                                        shading_elm[position] = parse_xml(r'<w:shd {} w:fill="e6e6e6"/>'.format(nsdecls('w')))
                                        shading_elm[position+99999] = parse_xml(r'<w:shd {} w:fill="e6e6e6"/>'.format(nsdecls('w')))
                                        table.rows[row*2].cells[col]._tc.get_or_add_tcPr().append(shading_elm[position])
                                        table.rows[row*2+1].cells[col]._tc.get_or_add_tcPr().append(shading_elm[position+99999])
                                    else:
                                        # just put all none alphabet chars in the row below:
                                        cell = table.cell(row *2 +1 , col)
                                        cell.text = secret

                    # make font smaller to fit into page and include hint at the end
                    docx_change_font_size(table, 4, 8)

                docx[player].add_paragraph(hint)


            # puzzles done

            # instructions for the parents of this player to hide the next stage
            hiding_instructions[player] += text_hide_next_round( round_num, total_rounds, player, hiding_place, language)

            if debug or player in include_plaintext_for_players:
                docx[player].add_paragraph('DEBUG PLAINTEXT: ' + message)


            docx[player].add_page_break()

    ### end of loop over rounds

    # special case of join_puzzle, then we need to add one more page for every player
    if join_method =='join_puzzle':
        # increase by one because it's one more puzzle after the loop
        msg_num = round_num +1
        for player in players_grade:
            if language == 'en':
                header = 'Message ' + str(msg_num  ) + ' for ' + player
                docx[player].add_heading(header)
                docx[player].add_paragraph(' ')
                docx[player].add_paragraph('This is a group puzzle. It can only be solved with all players together! Here is your share:')
                docx[player].add_paragraph(join_puzzle_messages[player])
            elif language == 'de':
                header = 'Nachricht ' + str(msg_num ) + ' für ' + player
                docx[player].add_heading(header)
                docx[player].add_paragraph(' ')
                docx[player].add_paragraph('Dies ist ein Gruppenrätsel. Es kann nur mit allen Spielern zusammen gelöst werden! Hier ist Dein Anteil:')
                docx[player].add_paragraph(join_puzzle_messages[player])


            if debug or player in include_plaintext_for_players:
                docx[player].add_paragraph('DEBUG PLAINTEXT: ' + join_puzzle_complete_message)

            docx[player].add_page_break()

    stego_saurus_header = {}
    if 'combination_lock_puzzle' in globals() and combination_lock_puzzle == 'stego_saurus':

        # stego_saurus does only 3 chars at the moment, brute force the 4th if more dials in lock 
        # 
        message = combination_lock_combination[0:3]
        ascii_saurus, hint = crypto_puzzles.stego_saurus(message, language, grade)

        log('encrypted_message:' + ascii_saurus)

        for player in players_grade:
            if language == 'en':
                stego_saurus_header[player] = 'Final Message for ' + player
                docx[player].add_heading(stego_saurus_header[player])
                docx[player].add_paragraph(' ')
                docx[player].add_paragraph('Here you might find the the first 3 digits from the combination for the treasure chest! If it has 4 dials, just try the last ;)')
            elif language == 'de':
                stego_saurus_header[player] = 'Letzte Nachricht für ' + player
                docx[player].add_heading(stego_saurus_header[player])
                docx[player].add_paragraph(' ')
                docx[player].add_paragraph('Hier kannst Du die ersten 3 Stellen der Kombination für das Zahlenschloß erfahren! Wenn das Schloß 4 Stellen hat, probier die letzten einfach durch ;)')

            paragraph= docx[player].add_paragraph()
            font = paragraph.add_run(ascii_saurus).font
            font.name = "FreeMono" # works on Linux, don't know about win (todo?)

            # make smaller for higher grades to make the content more difficult to spot 
            if grade < 5:
                font.size = Pt(7)
            elif grade < 7:
                font.size = Pt(5)
            else:
                font.size = Pt(4)
            # font size 2 is barely readable on my inkjet even with fine printing quality


            docx[player].add_paragraph(hint)
            docx[player].add_page_break()


    # puzzles done, now instructions for parents and stuff
    
    for player in players_grade:

        # if uv lamps are used, include page which explains them
        if uv_lamps_place:
            if language == 'en':
                title_uv ='Special info on ultraviolet lamps for ' + player
                docx[player].add_heading(title_uv)
                docx[player].add_paragraph(' ')
                docx[player].add_paragraph('The ultraviolet light will lead you to the next message!')
                docx[player].add_paragraph('If you do not have enough lamps for everybody: Youngest first!')
                docx[player].add_paragraph('If it is too bright to see the marks, make it darker.')
                if grade < 5:
                    docx[player].add_paragraph('Start searching on the floor with the lamp.')
                docx[player].add_page_break()

                # hide this page above and lamp(s)
                hiding_instructions[player] += ('Ultraviolet pocket lamps (ideally one per player) and note with title: ' + title_uv + ':::' + uv_lamps_place + '\n')
                hiding_instructions[player] += ('Use ultraviolet marker to make arrows on the floor from "' + uv_lamps_place + '" to ...' + ':::' + final_treasure_place + '\n')

                if join_method =='tokens':
                    # hiding instructions for token collector webpage
                    hiding_instructions[player] += ('The file tokens.txt and the script ' + token_script + ' after change of the variable next_place in the script to: "' + uv_lamps_place  + '"::: Webserver which can be reached by every player. Change the URL in the introduction (either before execution in euli.py or in all generated .docx).\n')

            if language == 'de':
                title_uv ='Spezial-Info UV-Lampen für ' + player
                docx[player].add_heading(title_uv)
                docx[player].add_paragraph(' ')
                docx[player].add_paragraph('Das UV-Licht führt Euch zum nächsten Hinweis!')
                docx[player].add_paragraph('Bei ungenügend Lampen: Die Jüngsten zuerst!')
                docx[player].add_paragraph('Eventuell ist es zu hell um etwas zu erkennen, dann einfach verdunkeln.')
                if grade < 5:
                    docx[player].add_paragraph('Fang auf dem Boden an zu suchen mit der Lampe.')

                docx[player].add_page_break()

                # hide this page above and lamp(s)
                hiding_instructions[player] += ('UV-Lampen (im Idealfall eine pro Spieler) und Zettel mit Titel: ' + title_uv + ':::' + uv_lamps_place + '\n')
                hiding_instructions[player] += ('UV-Stift Pfeile auf dem Boden von "' + uv_lamps_place + '" machen bis zu ... :::' + final_treasure_place + '\n')

                if join_method =='tokens':
                    # hiding instructions for token collector webpage
                    hiding_instructions[player] += ('Die Datei tokens.txt und das Script ' + token_script + ' mit Anpassung der Variablen next_place im Script auf: "' + uv_lamps_place  + '"::: Webserver, der für alle Spielern erreichbar ist. Die URL in der Einführung entsprechend anpassen (in euli.py oder in allen .docx).\n')

        # create list of what to hide where for the parents
        if language == 'en':
            docx[player].add_heading('Instructions to hide the parts for the treasure hunt of ' + player)
        elif language == 'de':
            docx[player].add_heading('Anleitung zum Verstecken der einzelnen Teile für die Schatzsuche von ' + player)

        docx[player].add_paragraph(parents_introduction[language])


        instructions = hiding_instructions[player].splitlines()
        # TODO: order of instructions is just by alphabetical sorting and starting sentences with 'A' if something should be on top, improve someday
        instructions.sort()

        if 'combination_lock_puzzle' in globals() and combination_lock_puzzle == 'stego_saurus':
            instructions.append(stego_saurus_header[player] + ':::' + final_treasure_place)
        else:
            if language == 'en':
                instructions.append('Key (or note with combination for the lock) for the treasure chest:::' + final_treasure_place)
            elif language == 'de':
                instructions.append('Schlüssel (oder Zettel mit der Kombination bei einem Zahlenschloss) für die Schatz-Truhe:::' + final_treasure_place)

        log("#####" + str(instructions) + "#####")

        # built table for instructions
        r = len(instructions)  # Number of rows 
        c = 2 # Number of columns you want

        # +1 for heading
        table = docx[player].add_table(rows=r+1, cols=c)

        # heading
        row = table.rows[0]
        if language == 'en':
            row.cells[0].text = 'Checkbox / What to hide'
            row.cells[1].text = 'Where to hide (in perspective of ' + player + ')'
        elif language == 'de':
            row.cells[0].text = 'Ankreuzfeld / Was verstecken'
            row.cells[1].text = 'Wo verstecken (aus Sicht von ' + player + ')'

        for y in range(r):
            what, where = instructions[y].split(':::')
            row = table.rows[y+1]
            row.cells[0].text = '❏ ' + what
            row.cells[1].text = where

        make_rows_bold(table.rows[0])

        if join_method =='tokens':
            # empty line for more space
            docx[player].add_paragraph(' ')
            paragraph = docx[player].add_paragraph('')
            if language == 'en':
                run = paragraph.add_run('Cut out the following secret tokens in the locations named above. Sticky tape works best.\n ✁ - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - ')
            elif language == 'de':
                run = paragraph.add_run('Die folgenden Geheim-Codes bitte einzeln ausschneiden und verstecken. Am besten mit Klebestreifen befestigen.\n ✁ - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - ')
            run.bold = True

            tokens = token_cut_outs[player].splitlines()
            log("Tokens:"+ ('    '.join(tokens)))
            docx[player].add_paragraph('   '.join(tokens))

        
        # Instructions how to incorporate own puzzles
        if language == 'en':
            docx[player].add_heading('Optional: Instructions how to incorporate own puzzles')
        elif language == 'de':
            docx[player].add_heading('Optional: Anleitung zum Einfügen eigener Rätsel')

        docx[player].add_paragraph(parents_introduction_hide_own_puzzles[language])


        # mark output filenames with _DEBUG to avoid using files which contain all plaintexts in real life ;)
        debug_insert = ''
        if debug:
            debug_insert = '_DEBUG'

        outfile = outdir + '/setup_instructions_euli_' + player + debug_insert + '.docx'
        docx[player].save(outfile) # Save document

    print("\nDone, wrote all files to " + outdir)

if __name__ == "__main__":
    main()


